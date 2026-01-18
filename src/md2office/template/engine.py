"""Template engine for variable injection and processing."""

from __future__ import annotations

import re
from typing import Any

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from jinja2 import BaseLoader, Environment, TemplateSyntaxError

from md2office.core.exceptions import TemplateError
from md2office.template.storage import TemplateStorage


class StringLoader(BaseLoader):
    """Jinja2 loader for string templates."""

    def get_source(self, environment: Environment, template: str):
        return template, None, lambda: True


class TemplateEngine:
    """Handles DOCX template loading and variable injection.

    Supports Jinja2-style variables ({{variable}}) in:
    - Document body paragraphs
    - Headers and footers
    - Table cells
    """

    # Pattern for finding {{variable}} placeholders
    VARIABLE_PATTERN = re.compile(r"\{\{([^}]+)\}\}")

    def __init__(self, storage: TemplateStorage | None = None) -> None:
        """Initialize template engine.

        Args:
            storage: Optional template storage instance.
        """
        self._storage = storage or TemplateStorage()
        self._jinja_env = Environment(loader=StringLoader())

    def load_template(self, name: str) -> Document:
        """Load a template document.

        Args:
            name: Template name.

        Returns:
            Word Document object.

        Raises:
            TemplateError: If template cannot be loaded.
        """
        try:
            path = self._storage.get_template_path(name)
            return Document(path)
        except Exception as e:
            raise TemplateError(f"Failed to load template: {e}") from e

    def inject_variables(
        self,
        document: Document,
        variables: dict[str, Any],
    ) -> Document:
        """Inject variables into a document.

        Replaces {{variable}} placeholders with values from the
        variables dictionary throughout the document.

        Args:
            document: Word document to modify.
            variables: Dictionary of variable names and values.

        Returns:
            Modified document.
        """
        # Process body paragraphs
        for paragraph in document.paragraphs:
            self._inject_in_paragraph(paragraph, variables)

        # Process tables
        for table in document.tables:
            self._inject_in_table(table, variables)

        # Process headers
        for section in document.sections:
            header = section.header
            if header:
                for paragraph in header.paragraphs:
                    self._inject_in_paragraph(paragraph, variables)
                for table in header.tables:
                    self._inject_in_table(table, variables)

            # Process footers
            footer = section.footer
            if footer:
                for paragraph in footer.paragraphs:
                    self._inject_in_paragraph(paragraph, variables)
                for table in footer.tables:
                    self._inject_in_table(table, variables)

        return document

    def _inject_in_paragraph(
        self,
        paragraph: Paragraph,
        variables: dict[str, Any],
    ) -> None:
        """Inject variables into a paragraph.

        Args:
            paragraph: Paragraph to process.
            variables: Variables dictionary.
        """
        # Check if paragraph contains any variables
        full_text = paragraph.text
        if "{{" not in full_text:
            return

        # Replace variables in the text
        new_text = self._replace_variables(full_text, variables)

        if new_text != full_text:
            # Clear existing runs and add new text
            # This preserves the first run's formatting
            if paragraph.runs:
                first_run = paragraph.runs[0]
                # Store formatting
                bold = first_run.bold
                italic = first_run.italic
                font_name = first_run.font.name
                font_size = first_run.font.size

                # Clear all runs
                for run in paragraph.runs:
                    run.text = ""

                # Set new text on first run
                first_run.text = new_text
                first_run.bold = bold
                first_run.italic = italic
                if font_name:
                    first_run.font.name = font_name
                if font_size:
                    first_run.font.size = font_size

    def _inject_in_table(
        self,
        table: Table,
        variables: dict[str, Any],
    ) -> None:
        """Inject variables into a table.

        Args:
            table: Table to process.
            variables: Variables dictionary.
        """
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._inject_in_paragraph(paragraph, variables)

    def _replace_variables(
        self,
        text: str,
        variables: dict[str, Any],
    ) -> str:
        """Replace variable placeholders in text.

        Args:
            text: Text containing {{variable}} placeholders.
            variables: Dictionary of variable values.

        Returns:
            Text with variables replaced.
        """
        def replace_match(match: re.Match) -> str:
            var_name = match.group(1).strip()
            if var_name in variables:
                return str(variables[var_name])
            # Leave unmatched variables as-is
            return match.group(0)

        return self.VARIABLE_PATTERN.sub(replace_match, text)

    def extract_variables(self, document: Document) -> set[str]:
        """Extract all variable names from a document.

        Args:
            document: Document to scan.

        Returns:
            Set of variable names found.
        """
        variables: set[str] = set()

        # Scan body
        for paragraph in document.paragraphs:
            variables.update(self._extract_from_text(paragraph.text))

        # Scan tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        variables.update(self._extract_from_text(paragraph.text))

        # Scan headers and footers
        for section in document.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    variables.update(self._extract_from_text(paragraph.text))
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    variables.update(self._extract_from_text(paragraph.text))

        return variables

    def _extract_from_text(self, text: str) -> set[str]:
        """Extract variable names from text.

        Args:
            text: Text to scan.

        Returns:
            Set of variable names found.
        """
        return {match.group(1).strip() for match in self.VARIABLE_PATTERN.finditer(text)}

    def process_template(
        self,
        name: str,
        variables: dict[str, Any] | None = None,
    ) -> Document:
        """Load a template and inject variables.

        Convenience method combining load and inject.

        Args:
            name: Template name.
            variables: Optional variables to inject.

        Returns:
            Processed document.
        """
        document = self.load_template(name)
        if variables:
            document = self.inject_variables(document, variables)
        return document

    def render_jinja_string(self, template_str: str, variables: dict[str, Any]) -> str:
        """Render a Jinja2 template string.

        For advanced templating in Markdown content.

        Args:
            template_str: Jinja2 template string.
            variables: Variables for rendering.

        Returns:
            Rendered string.

        Raises:
            TemplateError: If rendering fails.
        """
        try:
            template = self._jinja_env.from_string(template_str)
            return template.render(**variables)
        except TemplateSyntaxError as e:
            raise TemplateError(f"Template syntax error: {e}") from e
        except Exception as e:
            raise TemplateError(f"Template rendering failed: {e}") from e
