"""Generate default DOCX templates."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

from md2office.utils.helpers import ensure_dir


def create_default_template(output_path: Path | str | None = None) -> Path:
    """Create a default DOCX template with standard styles.

    Args:
        output_path: Path to save the template.
                    Defaults to templates/default.docx

    Returns:
        Path to the created template.
    """
    if output_path is None:
        output_path = Path("templates/default.docx")
    else:
        output_path = Path(output_path)

    ensure_dir(output_path.parent)

    doc = Document()

    # Configure document margins
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Define heading styles
    _configure_heading_styles(doc)

    # Define code styles
    _configure_code_styles(doc)

    # Define list styles
    _configure_list_styles(doc)

    # Define quote style
    _configure_quote_style(doc)

    # Add header with placeholder
    _configure_header(doc)

    # Add footer with page number
    _configure_footer(doc)

    # Save the template
    doc.save(output_path)

    return output_path


def _configure_heading_styles(doc: Document) -> None:
    """Configure heading styles."""
    styles = doc.styles

    # Heading configurations: (level, size, bold, color)
    heading_config = [
        (1, 24, True, "2F5496"),
        (2, 18, True, "2F5496"),
        (3, 14, True, "2F5496"),
        (4, 12, True, "2F5496"),
        (5, 11, True, "2F5496"),
        (6, 11, False, "2F5496"),
    ]

    for level, size, bold, color in heading_config:
        style_name = f"Heading {level}"
        try:
            style = styles[style_name]
        except KeyError:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if level == 1 else 6)
        style.paragraph_format.space_after = Pt(6)


def _configure_code_styles(doc: Document) -> None:
    """Configure code styles."""
    styles = doc.styles

    # Inline code character style
    try:
        code_char = styles["Code Char"]
    except KeyError:
        code_char = styles.add_style("Code Char", WD_STYLE_TYPE.CHARACTER)

    code_char.font.name = "Consolas"
    code_char.font.size = Pt(10)

    # Code block paragraph style
    try:
        code_block = styles["Code Block"]
    except KeyError:
        code_block = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)

    code_block.font.name = "Consolas"
    code_block.font.size = Pt(10)
    code_block.paragraph_format.left_indent = Inches(0.5)
    code_block.paragraph_format.space_before = Pt(6)
    code_block.paragraph_format.space_after = Pt(6)


def _configure_list_styles(doc: Document) -> None:
    """Configure list styles."""
    styles = doc.styles

    # List Bullet style
    try:
        list_bullet = styles["List Bullet"]
    except KeyError:
        list_bullet = styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)

    list_bullet.paragraph_format.left_indent = Inches(0.5)
    list_bullet.paragraph_format.space_after = Pt(3)

    # List Number style
    try:
        list_number = styles["List Number"]
    except KeyError:
        list_number = styles.add_style("List Number", WD_STYLE_TYPE.PARAGRAPH)

    list_number.paragraph_format.left_indent = Inches(0.5)
    list_number.paragraph_format.space_after = Pt(3)


def _configure_quote_style(doc: Document) -> None:
    """Configure blockquote style."""
    styles = doc.styles

    try:
        quote = styles["Quote"]
    except KeyError:
        quote = styles.add_style("Quote", WD_STYLE_TYPE.PARAGRAPH)

    quote.font.italic = True
    quote.font.color.rgb = RGBColor.from_string("666666")
    quote.paragraph_format.left_indent = Inches(0.5)
    quote.paragraph_format.space_before = Pt(6)
    quote.paragraph_format.space_after = Pt(6)


def _configure_header(doc: Document) -> None:
    """Configure document header."""
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "{{title}}"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _configure_footer(doc: Document) -> None:
    """Configure document footer with page number."""
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add page number field
    run = paragraph.add_run("Page ")
    _add_page_number_field(paragraph)


def _add_page_number_field(paragraph) -> None:
    """Add a page number field to a paragraph."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def create_professional_template(output_path: Path | str | None = None) -> Path:
    """Create a professional DOCX template with high-quality styles.

    This template uses modern typography and professional styling suitable
    for technical documentation, reports, and corporate documents.

    Args:
        output_path: Path to save the template.
                    Defaults to templates/professional.docx

    Returns:
        Path to the created template.
    """
    if output_path is None:
        output_path = Path("templates/professional.docx")
    else:
        output_path = Path(output_path)

    ensure_dir(output_path.parent)

    doc = Document()

    # Configure document margins (narrower for more content)
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    # Configure professional styles
    _configure_professional_normal(doc)
    _configure_professional_headings(doc)
    _configure_professional_code_styles(doc)
    _configure_professional_list_styles(doc)
    _configure_professional_quote_style(doc)
    _configure_professional_table_style(doc)

    # Add professional header/footer
    _configure_professional_header(doc)
    _configure_professional_footer(doc)

    # Save the template
    doc.save(output_path)

    return output_path


def _configure_professional_normal(doc: Document) -> None:
    """Configure the Normal style with professional typography."""
    style = doc.styles["Normal"]

    # Use a modern, readable font
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor.from_string("333333")

    # Good line spacing for readability
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(8)


def _configure_professional_headings(doc: Document) -> None:
    """Configure professional heading styles."""
    styles = doc.styles

    # Professional color scheme - dark blue gradient
    heading_config = [
        # (level, font_size, bold, color, space_before, space_after, font_name)
        (1, 28, True, "1B4F72", 24, 12, "Calibri Light"),
        (2, 22, True, "1F618D", 18, 10, "Calibri Light"),
        (3, 16, True, "2874A6", 14, 8, "Calibri"),
        (4, 13, True, "2E86C1", 12, 6, "Calibri"),
        (5, 12, True, "3498DB", 10, 4, "Calibri"),
        (6, 11, True, "5DADE2", 8, 4, "Calibri"),
    ]

    for level, size, bold, color, space_before, space_after, font_name in heading_config:
        style_name = f"Heading {level}"
        try:
            style = styles[style_name]
        except KeyError:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(space_after)
        style.paragraph_format.keep_with_next = True

        # Add bottom border for H1 and H2
        if level <= 2:
            _add_paragraph_border(style, "bottom", color, 8 if level == 1 else 4)


def _add_paragraph_border(style, position: str, color: str, size: int) -> None:
    """Add a border to a paragraph style."""
    pPr = style._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    border = OxmlElement(f"w:{position}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:color"), color)
    border.set(qn("w:space"), "4")
    pBdr.append(border)
    pPr.append(pBdr)


def _configure_professional_code_styles(doc: Document) -> None:
    """Configure professional code styles with background shading."""
    styles = doc.styles

    # Inline code character style
    try:
        code_char = styles["Code Char"]
    except KeyError:
        code_char = styles.add_style("Code Char", WD_STYLE_TYPE.CHARACTER)

    code_char.font.name = "Consolas"
    code_char.font.size = Pt(10)
    code_char.font.color.rgb = RGBColor.from_string("C7254E")

    # Code block paragraph style
    try:
        code_block = styles["Code Block"]
    except KeyError:
        code_block = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)

    code_block.font.name = "Consolas"
    code_block.font.size = Pt(9.5)
    code_block.font.color.rgb = RGBColor.from_string("333333")
    code_block.paragraph_format.left_indent = Inches(0.25)
    code_block.paragraph_format.right_indent = Inches(0.25)
    code_block.paragraph_format.space_before = Pt(8)
    code_block.paragraph_format.space_after = Pt(8)
    code_block.paragraph_format.line_spacing = 1.0

    # Add background shading
    _add_paragraph_shading(code_block, "F5F5F5")


def _add_paragraph_shading(style, color: str) -> None:
    """Add background shading to a paragraph style."""
    pPr = style._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    pPr.append(shd)


def _configure_professional_list_styles(doc: Document) -> None:
    """Configure professional list styles."""
    styles = doc.styles

    # List Bullet style - indentation handled by list builder
    try:
        list_bullet = styles["List Bullet"]
    except KeyError:
        list_bullet = styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)

    list_bullet.font.name = "Calibri"
    list_bullet.font.size = Pt(11)
    list_bullet.paragraph_format.space_after = Pt(3)
    list_bullet.paragraph_format.line_spacing = 1.15

    # List Number style - indentation handled by list builder
    try:
        list_number = styles["List Number"]
    except KeyError:
        list_number = styles.add_style("List Number", WD_STYLE_TYPE.PARAGRAPH)

    list_number.font.name = "Calibri"
    list_number.font.size = Pt(11)
    list_number.paragraph_format.space_after = Pt(3)
    list_number.paragraph_format.line_spacing = 1.15


def _configure_professional_quote_style(doc: Document) -> None:
    """Configure professional blockquote style."""
    styles = doc.styles

    try:
        quote = styles["Quote"]
    except KeyError:
        quote = styles.add_style("Quote", WD_STYLE_TYPE.PARAGRAPH)

    quote.font.name = "Georgia"
    quote.font.size = Pt(11)
    quote.font.italic = True
    quote.font.color.rgb = RGBColor.from_string("555555")
    quote.paragraph_format.left_indent = Inches(0.5)
    quote.paragraph_format.right_indent = Inches(0.5)
    quote.paragraph_format.space_before = Pt(10)
    quote.paragraph_format.space_after = Pt(10)
    quote.paragraph_format.line_spacing = 1.3

    # Add left border
    _add_paragraph_border(quote, "left", "CCCCCC", 16)


def _configure_professional_table_style(doc: Document) -> None:
    """Configure professional table style."""
    styles = doc.styles

    # Table Grid style usually exists, but we'll configure it
    try:
        table_style = styles["Table Grid"]
    except KeyError:
        table_style = styles.add_style("Table Grid", WD_STYLE_TYPE.TABLE)


def _configure_professional_header(doc: Document) -> None:
    """Configure professional document header."""
    section = doc.sections[0]
    header = section.header

    # Use a simple paragraph instead of table for better margin handling
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add title
    run = paragraph.add_run("{{title}}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("666666")

    # Add tab for right alignment
    run = paragraph.add_run("\t\t")

    # Add date
    run = paragraph.add_run("{{date}}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("666666")

    # Add bottom border
    _add_paragraph_border_simple(paragraph, "bottom", "CCCCCC", 4)


def _add_paragraph_border_simple(paragraph, position: str, color: str, size: int) -> None:
    """Add a border to a paragraph element directly."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    border = OxmlElement(f"w:{position}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:color"), color)
    border.set(qn("w:space"), "4")
    pBdr.append(border)
    pPr.append(pBdr)


def _remove_table_borders(table) -> None:
    """Remove all borders from a table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tbl_borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "nil")
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)


def _add_cell_bottom_border(cell, color: str) -> None:
    """Add a bottom border to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), color)
    tcBorders.append(bottom)
    tcPr.append(tcBorders)


def _configure_professional_footer(doc: Document) -> None:
    """Configure professional document footer with page number."""
    section = doc.sections[0]
    footer = section.footer

    # Use a simple paragraph instead of table for better margin handling
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add top border
    _add_paragraph_border_simple(paragraph, "top", "CCCCCC", 4)

    # Add author
    run = paragraph.add_run("{{author}}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("999999")

    # Add tabs for right alignment
    run = paragraph.add_run("\t\t")

    # Add page number
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("999999")
    _add_page_number_field_styled(paragraph, "999999")


def _add_cell_top_border(cell, color: str) -> None:
    """Add a top border to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:color"), color)
    tcBorders.append(top)
    tcPr.append(tcBorders)


def _add_page_number_field_styled(paragraph, color: str) -> None:
    """Add a styled page number field to a paragraph."""
    run = paragraph.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


if __name__ == "__main__":
    path = create_default_template()
    print(f"Created default template: {path}")
    path = create_professional_template()
    print(f"Created professional template: {path}")
