"""Main DOCX document builder."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import BinaryIO

from docx import Document
from docx.shared import Inches, Pt

from md2office.builder.admonition_builder import AdmonitionBuilder
from md2office.builder.list_builder import ListBuilder
from md2office.builder.style_mapper import StyleMapper
from md2office.builder.table_builder import TableBuilder
from md2office.core.config import StylesConfig
from md2office.core.exceptions import BuilderError
from md2office.parser.elements import (
    DocxAdmonition,
    DocxBlockquote,
    DocxCodeBlock,
    DocxElement,
    DocxHeading,
    DocxHorizontalRule,
    DocxImage,
    DocxList,
    DocxParagraph,
    DocxTable,
    Document as ASTDocument,
    TextSpan,
)


class DocxBuilder:
    """Builds Word documents from DOCX AST elements."""

    def __init__(
        self,
        template_path: Path | str | None = None,
        styles_config: StylesConfig | None = None,
    ) -> None:
        """Initialize the builder.

        Args:
            template_path: Optional path to a DOCX template.
            styles_config: Optional style configuration.
        """
        self._template_path = Path(template_path) if template_path else None
        self._styles_config = styles_config or StylesConfig()
        self._document: Document | None = None
        self._style_mapper: StyleMapper | None = None
        self._table_builder: TableBuilder | None = None
        self._list_builder: ListBuilder | None = None
        self._admonition_builder: AdmonitionBuilder | None = None

    def _init_document(self) -> None:
        """Initialize the Word document and helpers."""
        if self._template_path and self._template_path.exists():
            self._document = Document(self._template_path)
        else:
            self._document = Document()

        self._style_mapper = StyleMapper(self._styles_config, self._document)
        self._table_builder = TableBuilder(self._document, self._style_mapper)
        self._list_builder = ListBuilder(self._document, self._style_mapper)
        self._admonition_builder = AdmonitionBuilder(self._document, self._style_mapper)

    def build(self, elements: ASTDocument) -> Document:
        """Build a Word document from AST elements.

        Args:
            elements: List of DocxElement objects.

        Returns:
            Word Document object.

        Raises:
            BuilderError: If building fails.
        """
        try:
            self._init_document()
            assert self._document is not None

            for element in elements:
                self._build_element(element)

            return self._document
        except Exception as e:
            raise BuilderError(f"Failed to build document: {e}") from e

    def build_to_file(self, elements: ASTDocument, output_path: Path | str) -> Path:
        """Build and save a Word document to a file.

        Args:
            elements: List of DocxElement objects.
            output_path: Path to save the document.

        Returns:
            Path to the saved document.
        """
        document = self.build(elements)
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
        return output_path

    def build_to_bytes(self, elements: ASTDocument) -> bytes:
        """Build a Word document and return as bytes.

        Args:
            elements: List of DocxElement objects.

        Returns:
            Document as bytes.
        """
        document = self.build(elements)
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def build_to_stream(self, elements: ASTDocument, stream: BinaryIO) -> None:
        """Build a Word document and write to a stream.

        Args:
            elements: List of DocxElement objects.
            stream: Binary stream to write to.
        """
        document = self.build(elements)
        document.save(stream)

    def _build_element(self, element: DocxElement) -> None:
        """Build a single element and add to document.

        Args:
            element: DocxElement to build.
        """
        assert self._document is not None
        assert self._style_mapper is not None

        if isinstance(element, DocxHeading):
            self._build_heading(element)
        elif isinstance(element, DocxParagraph):
            self._build_paragraph(element)
        elif isinstance(element, DocxCodeBlock):
            self._build_code_block(element)
        elif isinstance(element, DocxBlockquote):
            self._build_blockquote(element)
        elif isinstance(element, DocxList):
            self._build_list(element)
        elif isinstance(element, DocxTable):
            self._build_table(element)
        elif isinstance(element, DocxImage):
            self._build_image(element)
        elif isinstance(element, DocxHorizontalRule):
            self._build_horizontal_rule()
        elif isinstance(element, DocxAdmonition):
            self._build_admonition(element)

    def _build_heading(self, heading: DocxHeading) -> None:
        """Build a heading element."""
        assert self._document is not None
        assert self._style_mapper is not None

        style = self._style_mapper.heading_style(heading.level)
        paragraph = self._document.add_paragraph()

        try:
            paragraph.style = style
        except KeyError:
            pass  # Style not found, use default

        self._add_text_spans(paragraph, heading.content)

    def _build_paragraph(self, para: DocxParagraph) -> None:
        """Build a paragraph element."""
        assert self._document is not None
        assert self._style_mapper is not None

        style = self._style_mapper.paragraph_style("normal")
        paragraph = self._document.add_paragraph()

        try:
            paragraph.style = style
        except KeyError:
            pass

        self._add_text_spans(paragraph, para.content)

    def _build_code_block(self, code_block: DocxCodeBlock) -> None:
        """Build a code block element."""
        assert self._document is not None
        assert self._style_mapper is not None

        style = self._style_mapper.code_style(block=True)
        paragraph = self._document.add_paragraph()

        try:
            paragraph.style = style
        except KeyError:
            pass

        run = paragraph.add_run(code_block.code)
        run.font.name = "Consolas"
        run.font.size = Pt(10)

    def _build_blockquote(self, blockquote: DocxBlockquote) -> None:
        """Build a blockquote element."""
        assert self._document is not None
        assert self._style_mapper is not None

        for child in blockquote.children:
            if isinstance(child, DocxParagraph):
                style = self._style_mapper.paragraph_style("quote")
                paragraph = self._document.add_paragraph()

                try:
                    paragraph.style = style
                except KeyError:
                    pass

                # Add indentation for quote
                paragraph.paragraph_format.left_indent = Inches(0.5)
                self._add_text_spans(paragraph, child.content)
            else:
                self._build_element(child)

    def _build_list(self, list_element: DocxList) -> None:
        """Build a list element."""
        assert self._list_builder is not None
        self._list_builder.build(list_element)

    def _build_table(self, table: DocxTable) -> None:
        """Build a table element."""
        assert self._table_builder is not None
        self._table_builder.build(table)

    def _build_image(self, image: DocxImage) -> None:
        """Build an image element."""
        assert self._document is not None

        paragraph = self._document.add_paragraph()

        # Try to add image if it's a local file
        image_path = Path(image.src)
        if image_path.exists():
            try:
                run = paragraph.add_run()
                run.add_picture(str(image_path), width=Inches(5))
            except Exception:
                # Fall back to placeholder text
                run = paragraph.add_run(f"[Image: {image.alt or image.src}]")
                run.italic = True
        else:
            # Placeholder for remote images
            run = paragraph.add_run(f"[Image: {image.alt or image.src}]")
            run.italic = True

        # Add caption if alt text provided
        if image.alt:
            caption = self._document.add_paragraph()
            caption_run = caption.add_run(image.alt)
            caption_run.italic = True
            caption_run.font.size = Pt(10)

    def _build_horizontal_rule(self) -> None:
        """Build a horizontal rule element."""
        assert self._document is not None

        # Add a paragraph with a bottom border to simulate HR
        paragraph = self._document.add_paragraph()
        paragraph.add_run("_" * 50)

    def _build_admonition(self, admonition: DocxAdmonition) -> None:
        """Build an admonition element."""
        assert self._admonition_builder is not None
        self._admonition_builder.build(admonition)

    def _add_text_spans(self, paragraph, spans: list[TextSpan]) -> None:
        """Add text spans to a paragraph.

        Args:
            paragraph: Word paragraph.
            spans: List of text spans to add.
        """
        for span in spans:
            run = paragraph.add_run(span.text)
            run.bold = span.bold
            run.italic = span.italic

            if span.code:
                run.font.name = "Consolas"
                run.font.size = Pt(10)

            if span.strikethrough:
                run.font.strike = True

            if span.link:
                # For now, just underline links
                # Full hyperlink support requires oxml manipulation
                run.underline = True
