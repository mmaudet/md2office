"""List construction for DOCX documents."""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from md2office.builder.style_mapper import StyleMapper
from md2office.parser.elements import DocxList, DocxListItem, TextSpan


class ListBuilder:
    """Builds Word lists from DocxList elements."""

    def __init__(self, document: Document, style_mapper: StyleMapper) -> None:
        """Initialize list builder.

        Args:
            document: Word document to add lists to.
            style_mapper: Style mapper for list styles.
        """
        self._document = document
        self._style_mapper = style_mapper

    def build(self, list_element: DocxList, level: int = 0) -> list[Paragraph]:
        """Build Word list paragraphs from a DocxList element.

        Args:
            list_element: DocxList element to convert.
            level: Nesting level for indentation.

        Returns:
            List of Word Paragraph objects.
        """
        paragraphs: list[Paragraph] = []
        style_name = self._style_mapper.list_style(list_element.ordered)

        for idx, item in enumerate(list_element.items):
            para = self._build_list_item(
                item,
                list_element.ordered,
                level,
                style_name,
                idx + list_element.start,
            )
            paragraphs.append(para)

            # Handle nested content
            for child in item.children:
                if isinstance(child, DocxList):
                    paragraphs.extend(self.build(child, level + 1))

        return paragraphs

    def _build_list_item(
        self,
        item: DocxListItem,
        ordered: bool,
        level: int,
        style_name: str,
        number: int,
    ) -> Paragraph:
        """Build a single list item paragraph.

        Args:
            item: List item element.
            ordered: Whether parent list is ordered.
            level: Nesting level.
            style_name: Word style name to use.
            number: Item number (for ordered lists).

        Returns:
            Word Paragraph object.
        """
        paragraph = self._document.add_paragraph()

        # Use Normal style to avoid doubled bullets from List Bullet/Number styles
        # We'll add manual prefix for consistent rendering
        try:
            paragraph.style = "Normal"
        except KeyError:
            pass

        # Force LEFT alignment
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Set indentation based on level
        # Level 0: 0.35 inch (within margins, room for bullet)
        # Level 1+: additional 0.35 inch per level
        base_indent = Inches(0.35)
        level_indent = Inches(0.35 * level)
        paragraph.paragraph_format.left_indent = base_indent + level_indent
        paragraph.paragraph_format.space_after = Pt(3)

        # Add bullet/number prefix
        prefix = f"{number}. " if ordered else "- "
        run = paragraph.add_run(prefix)
        run.bold = False

        # Add content with formatting
        self._add_text_spans(paragraph, item.content)

        return paragraph

    def _add_text_spans(self, paragraph: Paragraph, spans: list[TextSpan]) -> None:
        """Add text spans to a paragraph.

        Args:
            paragraph: Word paragraph.
            spans: List of text spans to add.
        """
        for span in spans:
            if span.link:
                self._add_hyperlink(paragraph, span)
            else:
                run = paragraph.add_run(span.text)
                run.bold = span.bold
                run.italic = span.italic

                if span.code:
                    run.font.name = "Liberation Mono"
                    run.font.size = Pt(9)

                if span.strikethrough:
                    run.font.strike = True

    def _add_hyperlink(self, paragraph: Paragraph, span: TextSpan) -> None:
        """Add a clickable hyperlink to a paragraph.

        Handles both external links (http://, https://, etc.) and
        internal links (starting with #) that point to bookmarks.

        Args:
            paragraph: Word paragraph.
            span: TextSpan with link URL.
        """
        # Create the hyperlink element
        hyperlink = OxmlElement("w:hyperlink")

        # Check if this is an internal link (bookmark reference)
        if span.link and span.link.startswith("#"):
            # Internal link - use w:anchor attribute
            anchor_name = span.link[1:]  # Remove the leading #
            hyperlink.set(qn("w:anchor"), anchor_name)
        else:
            # External link - use r:id relationship
            part = paragraph.part
            r_id = part.relate_to(
                span.link,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True,
            )
            hyperlink.set(qn("r:id"), r_id)

        # Create the run element inside the hyperlink
        new_run = OxmlElement("w:r")

        # Add run properties
        rPr = OxmlElement("w:rPr")

        # Bold
        if span.bold:
            bold_elem = OxmlElement("w:b")
            rPr.append(bold_elem)

        # Italic
        if span.italic:
            italic_elem = OxmlElement("w:i")
            rPr.append(italic_elem)

        # Underline (standard for hyperlinks)
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        rPr.append(underline)

        # Blue color (standard for hyperlinks)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        rPr.append(color)

        # Code formatting
        if span.code:
            font = OxmlElement("w:rFonts")
            font.set(qn("w:ascii"), "Liberation Mono")
            font.set(qn("w:hAnsi"), "Liberation Mono")
            rPr.append(font)
            size = OxmlElement("w:sz")
            size.set(qn("w:val"), "18")  # 9pt = 18 half-points
            rPr.append(size)

        # Strikethrough
        if span.strikethrough:
            strike = OxmlElement("w:strike")
            rPr.append(strike)

        new_run.append(rPr)

        # Add the text
        text_elem = OxmlElement("w:t")
        text_elem.text = span.text
        new_run.append(text_elem)

        hyperlink.append(new_run)

        # Append hyperlink to paragraph
        paragraph._p.append(hyperlink)

