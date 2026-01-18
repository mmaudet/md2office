"""Table construction for DOCX documents."""

from __future__ import annotations

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import Table, _Cell

from md2office.builder.style_mapper import StyleMapper
from md2office.parser.elements import DocxTable, TextSpan


class TableBuilder:
    """Builds Word tables from DocxTable elements."""

    def __init__(self, document: Document, style_mapper: StyleMapper) -> None:
        """Initialize table builder.

        Args:
            document: Word document to add tables to.
            style_mapper: Style mapper for table styles.
        """
        self._document = document
        self._style_mapper = style_mapper

    def build(self, table_element: DocxTable) -> Table:
        """Build a Word table from a DocxTable element.

        Args:
            table_element: DocxTable element to convert.

        Returns:
            Word Table object.
        """
        if not table_element.rows:
            return self._document.add_table(rows=0, cols=0)

        num_rows = len(table_element.rows)
        num_cols = max(len(row.cells) for row in table_element.rows)

        # Create table
        table = self._document.add_table(rows=num_rows, cols=num_cols)

        # Apply table style
        try:
            table.style = self._style_mapper.table_style()
        except KeyError:
            pass  # Style not found, use default

        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Get table formatting config
        config = self._style_mapper.table_config()

        # Track merged cells to skip filling
        merged_cells: set[tuple[int, int]] = set()

        # First pass: identify cells to merge and perform merges
        self._process_cell_merges(table_element, table, merged_cells)

        # Second pass: fill table cells
        for row_idx, row in enumerate(table_element.rows):
            is_header_row = row_idx == 0 and table_element.has_header
            word_row = table.rows[row_idx]

            # Set minimum row height for vertical centering to be visible
            # 400 twips = 20pt minimum height
            self._set_row_height(word_row, 400)

            for col_idx, cell in enumerate(row.cells):
                if col_idx >= num_cols:
                    break

                # Skip cells that are part of a merged region (not the origin)
                if (row_idx, col_idx) in merged_cells:
                    continue

                # Skip cells marked for merge_up or merge_left
                if cell.merge_up or cell.merge_left:
                    continue

                table_cell = word_row.cells[col_idx]
                self._fill_cell(table_cell, cell.content, is_header_row, config)

            # Apply alternating row colors
            if not is_header_row and config.get("alternating_rows", False):
                if row_idx % 2 == 0:  # Even rows (0-indexed, so 0, 2, 4...)
                    self._set_row_bg_color(word_row, config.get("alt_row_bg", "D9E2F3"))

        return table

    def _process_cell_merges(
        self,
        table_element: DocxTable,
        table: Table,
        merged_cells: set[tuple[int, int]],
    ) -> None:
        """Process cell merges (rowspan and colspan).

        Args:
            table_element: Source table element.
            table: Word table to apply merges to.
            merged_cells: Set to track merged cell positions.
        """
        num_rows = len(table_element.rows)
        num_cols = max(len(row.cells) for row in table_element.rows) if table_element.rows else 0

        # Process vertical merges (cells with merge_up=True)
        for col_idx in range(num_cols):
            row_idx = 0
            while row_idx < num_rows:
                row = table_element.rows[row_idx]
                if col_idx >= len(row.cells):
                    row_idx += 1
                    continue

                cell = row.cells[col_idx]

                # If not a merge_up cell, check how many merge_up cells follow
                if not cell.merge_up:
                    merge_count = 0
                    check_row = row_idx + 1
                    while check_row < num_rows:
                        check_row_cells = table_element.rows[check_row].cells
                        if col_idx >= len(check_row_cells):
                            break
                        if check_row_cells[col_idx].merge_up:
                            merge_count += 1
                            merged_cells.add((check_row, col_idx))
                            check_row += 1
                        else:
                            break

                    # Perform vertical merge if needed
                    if merge_count > 0:
                        start_cell = table.rows[row_idx].cells[col_idx]
                        end_cell = table.rows[row_idx + merge_count].cells[col_idx]
                        start_cell.merge(end_cell)

                    row_idx = check_row
                else:
                    row_idx += 1

        # Process horizontal merges (cells with merge_left=True)
        for row_idx in range(num_rows):
            row = table_element.rows[row_idx]
            col_idx = 0
            while col_idx < len(row.cells):
                cell = row.cells[col_idx]

                # If not a merge_left cell, check how many merge_left cells follow
                if not cell.merge_left:
                    merge_count = 0
                    check_col = col_idx + 1
                    while check_col < len(row.cells):
                        if row.cells[check_col].merge_left:
                            merge_count += 1
                            merged_cells.add((row_idx, check_col))
                            check_col += 1
                        else:
                            break

                    # Perform horizontal merge if needed
                    if merge_count > 0:
                        start_cell = table.rows[row_idx].cells[col_idx]
                        end_cell = table.rows[row_idx].cells[col_idx + merge_count]
                        start_cell.merge(end_cell)

                    col_idx = check_col
                else:
                    col_idx += 1

    def _fill_cell(
        self,
        cell: _Cell,
        content: list[TextSpan],
        is_header: bool,
        config: dict,
    ) -> None:
        """Fill a table cell with content.

        Args:
            cell: Word table cell.
            content: List of text spans.
            is_header: Whether this is a header cell.
            config: Table formatting configuration.
        """
        # Clear existing content
        cell.text = ""
        paragraph = cell.paragraphs[0]

        # Set vertical centering for all cells
        self._set_cell_vertical_alignment(cell, "center")

        # Remove paragraph spacing that could interfere with vertical centering
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

        # Add content with formatting
        for span in content:
            if span.link:
                # Create a clickable hyperlink
                self._add_hyperlink(paragraph, span, is_header)
            else:
                run = paragraph.add_run(span.text)
                run.bold = span.bold or is_header
                run.italic = span.italic
                if span.code:
                    run.font.name = "Consolas"
                    run.font.size = Pt(10)

        # Header cell formatting
        if is_header:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_bg = config.get("header_bg", "4472C4")
            header_text = config.get("header_text", "FFFFFF")
            self._set_cell_bg_color(cell, header_bg)
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor.from_string(header_text)

    def _set_cell_bg_color(self, cell: _Cell, color: str) -> None:
        """Set background color of a cell.

        Args:
            cell: Word table cell.
            color: Hex color string (without #).
        """
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), color)
        tc_pr.append(shd)

    def _set_row_bg_color(self, row, color: str) -> None:
        """Set background color for all cells in a row.

        Args:
            row: Word table row.
            color: Hex color string (without #).
        """
        for cell in row.cells:
            self._set_cell_bg_color(cell, color)

    def _set_cell_vertical_alignment(self, cell: _Cell, alignment: str) -> None:
        """Set vertical alignment of a cell using oxml.

        Args:
            cell: Word table cell.
            alignment: Alignment value ('top', 'center', 'bottom').
        """
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        # Remove any existing vAlign element
        for existing in tc_pr.findall(qn("w:vAlign")):
            tc_pr.remove(existing)
        # Add new vAlign element
        v_align = OxmlElement("w:vAlign")
        v_align.set(qn("w:val"), alignment)
        tc_pr.append(v_align)

    def _set_cell_margins(self, cell: _Cell, top: int = 0, bottom: int = 0, left: int = 0, right: int = 0) -> None:
        """Set cell margins/padding using oxml.

        Args:
            cell: Word table cell.
            top: Top margin in twips.
            bottom: Bottom margin in twips.
            left: Left margin in twips.
            right: Right margin in twips.
        """
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        # Remove any existing tcMar element
        for existing in tc_pr.findall(qn("w:tcMar")):
            tc_pr.remove(existing)
        # Add new tcMar element
        tc_mar = OxmlElement("w:tcMar")
        if top:
            top_elem = OxmlElement("w:top")
            top_elem.set(qn("w:w"), str(top))
            top_elem.set(qn("w:type"), "dxa")
            tc_mar.append(top_elem)
        if bottom:
            bottom_elem = OxmlElement("w:bottom")
            bottom_elem.set(qn("w:w"), str(bottom))
            bottom_elem.set(qn("w:type"), "dxa")
            tc_mar.append(bottom_elem)
        if left:
            left_elem = OxmlElement("w:left")
            left_elem.set(qn("w:w"), str(left))
            left_elem.set(qn("w:type"), "dxa")
            tc_mar.append(left_elem)
        if right:
            right_elem = OxmlElement("w:right")
            right_elem.set(qn("w:w"), str(right))
            right_elem.set(qn("w:type"), "dxa")
            tc_mar.append(right_elem)
        tc_pr.append(tc_mar)

    def _set_row_height(self, row, height_twips: int, exact: bool = False) -> None:
        """Set row height.

        Args:
            row: Word table row.
            height_twips: Height in twips (1/20 of a point).
            exact: If True, use exact height; otherwise use minimum height.
        """
        tr = row._tr
        tr_pr = tr.get_or_add_trPr()
        # Remove any existing trHeight element
        for existing in tr_pr.findall(qn("w:trHeight")):
            tr_pr.remove(existing)
        # Add new trHeight element
        tr_height = OxmlElement("w:trHeight")
        tr_height.set(qn("w:val"), str(height_twips))
        tr_height.set(qn("w:hRule"), "exact" if exact else "atLeast")
        tr_pr.append(tr_height)

    def _add_hyperlink(self, paragraph, span: TextSpan, is_header: bool = False) -> None:
        """Add a clickable hyperlink to a paragraph.

        Args:
            paragraph: Word paragraph.
            span: TextSpan with link URL.
            is_header: Whether this is a header cell.
        """
        # Get the document part to add the relationship
        part = paragraph.part
        r_id = part.relate_to(
            span.link,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )

        # Create the hyperlink element
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        # Create the run element inside the hyperlink
        new_run = OxmlElement("w:r")

        # Add run properties
        rPr = OxmlElement("w:rPr")

        # Bold (for header or if span is bold)
        if span.bold or is_header:
            bold_elem = OxmlElement("w:b")
            rPr.append(bold_elem)

        # Italic
        if span.italic:
            italic_elem = OxmlElement("w:i")
            rPr.append(italic_elem)

        # Underline (standard for hyperlinks)
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        rPr.append(underline)

        # Blue color (standard for hyperlinks)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        rPr.append(color)

        # Code formatting
        if span.code:
            font = OxmlElement("w:rFonts")
            font.set(qn("w:ascii"), "Consolas")
            font.set(qn("w:hAnsi"), "Consolas")
            rPr.append(font)
            size = OxmlElement("w:sz")
            size.set(qn("w:val"), "20")  # 10pt = 20 half-points
            rPr.append(size)

        # Strikethrough
        if span.strikethrough:
            strike = OxmlElement("w:strike")
            rPr.append(strike)

        new_run.append(rPr)

        # Add the text
        text_elem = OxmlElement("w:t")
        text_elem.text = span.text
        new_run.append(text_elem)

        hyperlink.append(new_run)

        # Append hyperlink to paragraph
        paragraph._p.append(hyperlink)
