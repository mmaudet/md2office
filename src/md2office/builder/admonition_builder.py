"""Admonition/callout construction for DOCX documents."""

from __future__ import annotations

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table

from md2office.builder.style_mapper import StyleMapper
from md2office.parser.elements import DocxAdmonition, DocxParagraph, TextSpan


class AdmonitionBuilder:
    """Builds Word admonitions (callouts) as styled tables."""

    # Admonition icons - using simple text symbols for better compatibility
    ICONS = {
        "NOTE": "i",  # Information
        "TIP": "?",  # Tip/hint (will be styled as lightbulb visually)
        "IMPORTANT": "!",  # Important
        "WARNING": "!",  # Warning
        "CAUTION": "X",  # Caution/danger
    }

    def __init__(self, document: Document, style_mapper: StyleMapper) -> None:
        """Initialize admonition builder.

        Args:
            document: Word document to add admonitions to.
            style_mapper: Style mapper for admonition configuration.
        """
        self._document = document
        self._style_mapper = style_mapper

    def build(self, admonition: DocxAdmonition) -> Table:
        """Build a Word table representing an admonition.

        Args:
            admonition: DocxAdmonition element to convert.

        Returns:
            Word Table object styled as an admonition.
        """
        config = self._style_mapper.admonition_config(admonition.admonition_type)

        # Create a single-row, two-column table
        table = self._document.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False

        # Set table to full page width (6.5 inches for standard margins)
        self._set_table_full_width(table)

        # Set column widths (icon narrow, content fills rest)
        table.columns[0].width = Inches(0.6)  # Icon column
        table.columns[1].width = Inches(5.9)  # Content column (6.5 - 0.6)

        # Get the row and cells
        row = table.rows[0]
        icon_cell = row.cells[0]
        content_cell = row.cells[1]

        # Set minimum row height for vertical centering to be visible
        # 720 twips = 36pt = ~0.5 inch minimum height
        self._set_row_height(row, 720)

        # Style the icon cell
        self._build_icon_cell(icon_cell, admonition.admonition_type, config)

        # Style the content cell
        self._build_content_cell(content_cell, admonition, config)

        # Apply background color to both cells
        bg_color = config.get("bg", "DDF4FF")
        self._set_cell_bg_color(icon_cell, bg_color)
        self._set_cell_bg_color(content_cell, bg_color)

        # Apply border color
        border_color = config.get("color", "0969DA")
        self._set_table_borders(table, border_color)

        return table

    def _build_icon_cell(self, cell, admonition_type: str, config: dict) -> None:
        """Build the icon cell.

        Args:
            cell: Word table cell for the icon.
            admonition_type: Type of admonition.
            config: Admonition configuration.
        """
        cell.text = ""
        self._set_cell_vertical_alignment(cell, "center")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

        # Add icon with bold styling for better visibility
        icon = self.ICONS.get(admonition_type, "i")
        run = paragraph.add_run(icon)
        run.font.name = "Arial"
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(config.get("color", "0969DA"))

    def _build_content_cell(self, cell, admonition: DocxAdmonition, config: dict) -> None:
        """Build the content cell.

        Args:
            cell: Word table cell for content.
            admonition: Admonition element.
            config: Admonition configuration.
        """
        cell.text = ""
        self._set_cell_vertical_alignment(cell, "center")

        # Add content from children (no separate title)
        first_para = True
        for child in admonition.children:
            if isinstance(child, DocxParagraph):
                if first_para:
                    para = cell.paragraphs[0]
                    first_para = False
                else:
                    para = cell.add_paragraph()
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                self._add_text_spans(para, child.content)
                # Set text color for content
                for run in para.runs:
                    run.font.color.rgb = RGBColor.from_string(config.get("color", "0969DA"))

    def _add_text_spans(self, paragraph, spans: list[TextSpan]) -> None:
        """Add text spans to a paragraph.

        Args:
            paragraph: Word paragraph.
            spans: List of text spans to add.
        """
        for span in spans:
            run = paragraph.add_run(span.text)
            run.bold = span.bold
            run.italic = span.italic

            if span.code:
                run.font.name = "Consolas"
                run.font.size = Pt(10)

            if span.strikethrough:
                run.font.strike = True

    def _set_cell_bg_color(self, cell, color: str) -> None:
        """Set background color of a cell.

        Args:
            cell: Word table cell.
            color: Hex color string (without #).
        """
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), color)
        tc_pr.append(shd)

    def _set_table_borders(self, table: Table, color: str) -> None:
        """Set colored left border on the table.

        Args:
            table: Word table.
            color: Hex color string (without #).
        """
        tbl = table._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

        tbl_borders = OxmlElement("w:tblBorders")

        # Only add left border (callout style)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "24")  # 3pt border
        left.set(qn("w:color"), color)
        tbl_borders.append(left)

        # Remove other borders
        for border_name in ["top", "right", "bottom", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "nil")
            tbl_borders.append(border)

        tbl_pr.append(tbl_borders)
        if tbl.tblPr is None:
            tbl.insert(0, tbl_pr)

    def _set_table_full_width(self, table: Table) -> None:
        """Set table to full page width (100%).

        Args:
            table: Word table to set width.
        """
        tbl = table._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

        # Set table width to 100%
        tbl_w = OxmlElement("w:tblW")
        tbl_w.set(qn("w:w"), "5000")  # 5000 = 100% in fifths of a percent
        tbl_w.set(qn("w:type"), "pct")
        tbl_pr.append(tbl_w)

        if tbl.tblPr is None:
            tbl.insert(0, tbl_pr)

    def _set_cell_vertical_alignment(self, cell, alignment: str) -> None:
        """Set vertical alignment of a cell using oxml.

        Args:
            cell: Word table cell.
            alignment: Alignment value ('top', 'center', 'bottom').
        """
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        # Remove any existing vAlign element
        for existing in tc_pr.findall(qn("w:vAlign")):
            tc_pr.remove(existing)
        # Add new vAlign element
        v_align = OxmlElement("w:vAlign")
        v_align.set(qn("w:val"), alignment)
        tc_pr.append(v_align)

    def _set_row_height(self, row, height_twips: int) -> None:
        """Set minimum row height.

        Args:
            row: Word table row.
            height_twips: Height in twips (1/20 of a point).
        """
        tr = row._tr
        tr_pr = tr.get_or_add_trPr()
        # Remove any existing trHeight element
        for existing in tr_pr.findall(qn("w:trHeight")):
            tr_pr.remove(existing)
        # Add new trHeight element with atLeast rule
        tr_height = OxmlElement("w:trHeight")
        tr_height.set(qn("w:val"), str(height_twips))
        tr_height.set(qn("w:hRule"), "atLeast")
        tr_pr.append(tr_height)

    def _set_cell_margins(self, cell, top=None, bottom=None, left=None, right=None) -> None:
        """Set cell margins/padding using oxml.

        Args:
            cell: Word table cell.
            top: Top margin (Pt value).
            bottom: Bottom margin (Pt value).
            left: Left margin (Pt value).
            right: Right margin (Pt value).
        """
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        # Remove any existing tcMar element
        for existing in tc_pr.findall(qn("w:tcMar")):
            tc_pr.remove(existing)
        # Add new tcMar element
        tc_mar = OxmlElement("w:tcMar")
        if top:
            top_elem = OxmlElement("w:top")
            top_elem.set(qn("w:w"), str(int(top)))
            top_elem.set(qn("w:type"), "dxa")
            tc_mar.append(top_elem)
        if bottom:
            bottom_elem = OxmlElement("w:bottom")
            bottom_elem.set(qn("w:w"), str(int(bottom)))
            bottom_elem.set(qn("w:type"), "dxa")
            tc_mar.append(bottom_elem)
        if left:
            left_elem = OxmlElement("w:left")
            left_elem.set(qn("w:w"), str(int(left)))
            left_elem.set(qn("w:type"), "dxa")
            tc_mar.append(left_elem)
        if right:
            right_elem = OxmlElement("w:right")
            right_elem.set(qn("w:w"), str(int(right)))
            right_elem.set(qn("w:type"), "dxa")
            tc_mar.append(right_elem)
        tc_pr.append(tc_mar)
